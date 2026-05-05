"""Run python-docx native ingest on a .docx file.

Format-aware ingest — DOCX has no page concept (it's a flow document), so
output is one JSON + one Markdown for the whole document. Provenance for
DOCX-source citations is `(paragraph_id, cell_id)` per spec §10.5, NOT
`(page, bbox)` like PDFs.

Outputs:
  <out>/document.json   — structured extraction. Each block has:
                          {kind, paragraph_id, style, text}
                          plus, when the paragraph contains tab-separated
                          cells (line-item rows), `cells: [...]`. Table
                          blocks instead carry {kind: 'table_cell',
                          table_cell_id, row, col, text}.
  <out>/document.md     — Markdown rendering, one paragraph per line.

paragraph_id is the Word-native `w14:paraId` attribute on each <w:p>
element. Word writes it once per paragraph and preserves it across edits
to that paragraph; it is the on-disk equivalent of bbox for the DOCX
path. Coverage on the AZ budget bill (SB 1735) was 100% on probe — every
paragraph in the bill has one.

Why python-docx and not Docling: Docling proved unworkable on Windows
(see Phase 0 plan "Pivot — 2026-05-05"). python-docx is a thin XML reader
with no native bridge or model dependency, and the AZ Legislature's
`.docx` files are clean structured XML to begin with.
"""

import argparse
import json
import sys
from pathlib import Path


# Word 2010 namespace where paraId lives. python-docx exposes the underlying
# lxml element via paragraph._element; the attribute key is fully namespaced.
PARA_ID_ATTR = "{http://schemas.microsoft.com/office/word/2010/wordml}paraId"


def write_dry_run(docx: Path, out: Path) -> None:
    """Stub for testing — writes a minimal valid output without parsing."""
    out.mkdir(parents=True, exist_ok=True)
    (out / "document.json").write_text(
        json.dumps(
            {
                "extractor": "python-docx-dry-run",
                "source_docx": str(docx),
                "blocks": [],
            },
            indent=2,
        )
    )
    (out / "document.md").write_text("# Dry-run\n")


def _paragraph_id(para_element, fallback_index: int) -> str:
    """Stable id for a paragraph. Prefers Word's paraId; falls back to the
    1-based source-order index (recorded so the mapping is reproducible).
    """
    raw = para_element.attrib.get(PARA_ID_ATTR)
    if raw:
        return f"p:{raw}"
    return f"p:idx-{fallback_index}"


def _split_tab_cells(text: str) -> list[str] | None:
    """Return tab-split cells if the paragraph looks like a line-item row.

    AZ Legislature budget bills encode line-item rows as paragraphs with
    tab-separated columns (e.g. `JOBS\\t11,005,600`). Returning the cell
    list lets downstream chunking treat the row as a structured record
    even though it lives in a single <w:p> element.

    Returns None when there are no tabs — caller should not add a `cells`
    field in that case.
    """
    if "\t" not in text:
        return None
    return [c.strip() for c in text.split("\t")]


def run_docx_ingest(docx: Path, out: Path) -> None:
    """Real path. Walks the document body in source order and emits one
    block per paragraph + one block per table cell.

    Source order matters — citations need to render in the same order the
    bill reads. python-docx's `Document(docx).paragraphs` and `.tables`
    are flat lists that don't preserve interleave; we walk the document
    body's XML children directly.
    """
    # Local import keeps the dry-run path light for tests.
    from docx import Document
    from docx.oxml.ns import qn

    out.mkdir(parents=True, exist_ok=True)
    doc = Document(str(docx))
    body = doc.element.body

    blocks: list[dict] = []
    para_index = 0
    table_index = 0

    # Walk the body's direct children in source order. <w:p> = paragraph,
    # <w:tbl> = table; anything else (sectPr, bookmarkStart/End, etc.) we
    # skip silently — they're either structural noise or already captured
    # as paragraph properties.
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            para_index += 1
            # Reconstruct text from runs to preserve tabs (which python-docx's
            # Paragraph.text drops the tab character on some platforms — we
            # need it for line-item cell splitting).
            text_parts: list[str] = []
            for run in child.iter(qn("w:r")):
                for node in run.iter():
                    if node.tag == qn("w:t"):
                        text_parts.append(node.text or "")
                    elif node.tag == qn("w:tab"):
                        text_parts.append("\t")
                    elif node.tag == qn("w:br"):
                        text_parts.append("\n")
            text = "".join(text_parts)

            # Look up the style name via python-docx's high-level API since
            # the raw <w:pStyle val="..."> only carries the style id, not
            # the human-readable name we want for rubric scoring. When a
            # paragraph carries no <w:pStyle>, Word treats it as "Normal"
            # by default — we record that explicitly rather than None so
            # downstream consumers can grep on style without dealing with
            # null fallback logic.
            pstyle = child.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            style_id = pstyle.get(qn("w:val")) if pstyle is not None else None
            if style_id is None:
                style_name = "Normal"
            else:
                try:
                    style_name = doc.styles.element.get_by_id(style_id).name_val
                except Exception:
                    style_name = style_id  # fall back to the raw id

            block: dict = {
                "kind": "paragraph",
                "paragraph_id": _paragraph_id(child, para_index),
                "style": style_name,
                "text": text,
            }
            cells = _split_tab_cells(text)
            if cells is not None:
                block["cells"] = cells
            blocks.append(block)

        elif tag == qn("w:tbl"):
            table_index += 1
            # Walk rows / cells in source order. table_cell_id is
            # deterministic from position so re-runs produce identical ids.
            for row_index, row in enumerate(child.iter(qn("w:tr")), start=1):
                # Direct children only — nested tables would otherwise
                # collapse into the parent's cell list.
                for col_index, cell in enumerate(
                    [c for c in row.iter(qn("w:tc")) if c.getparent() is row],
                    start=1,
                ):
                    cell_text_parts: list[str] = []
                    for t in cell.iter(qn("w:t")):
                        cell_text_parts.append(t.text or "")
                    cell_text = "".join(cell_text_parts).strip()
                    blocks.append(
                        {
                            "kind": "table_cell",
                            "table_cell_id": (
                                f"t{table_index}-r{row_index}-c{col_index}"
                            ),
                            "row": row_index,
                            "col": col_index,
                            "text": cell_text,
                        }
                    )

    (out / "document.json").write_text(
        json.dumps(
            {
                "extractor": "python-docx",
                "source_docx": str(docx),
                "blocks": blocks,
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    # Markdown rendering: one block per line. Tab-separated cells become
    # `cell1 | cell2 | ...` so the .md is greppable by analysts. Table
    # cells render as `[t1-r1-c1] text`.
    md_lines: list[str] = []
    for b in blocks:
        if b["kind"] == "paragraph":
            text = b["text"].strip()
            if not text:
                continue
            if "cells" in b:
                md_lines.append(" | ".join(b["cells"]))
            else:
                md_lines.append(text)
        elif b["kind"] == "table_cell":
            text = b["text"].strip()
            if text:
                md_lines.append(f"[{b['table_cell_id']}] {text}")
    (out / "document.md").write_text("\n\n".join(md_lines) + "\n", encoding="utf-8")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Extract a .docx file into per-paragraph JSON + Markdown."
    )
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="skip real extraction (test mode)",
    )
    args = parser.parse_args(argv)

    if not args.docx.exists():
        print(f"DOCX not found: {args.docx}", file=sys.stderr)
        return 2

    if args.dry_run:
        write_dry_run(args.docx, args.out)
    else:
        run_docx_ingest(args.docx, args.out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
