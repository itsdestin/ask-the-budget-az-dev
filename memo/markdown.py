"""Markdown -> the memo's body.

Ported from `harness/documents.py`'s renderer, remapped onto JLBC's
styles. The regexes and the classification ORDER are carried over
unchanged: each encodes a fix for a real defect, and the ordering one is
load-bearing (see `render_body`).

THE RULE THAT MATTERS, and it is unchanged: anything unrecognized becomes
a plain paragraph, verbatim. Never a silent drop. An analyst who receives
a memo with a section quietly missing has no way to know it happened, and
that is a far worse failure than a blockquote rendering as ordinary text.

The recognized subset is deliberately small — it is the inverse of what
`primer/docx_to_md.py` emits, which is also what the system prompt tells
the model to write.
"""
from __future__ import annotations

import re

from docx.document import Document as DocumentT

from memo.style import BULLET_INDENT

_HEADING_RE = re.compile(r"^(#{1,6})\s+(\S.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(\S.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
# A separator row: only dashes, colons, pipes and spaces, and at least one
# dash. Requiring the PRECEDING line to be a row too (see the loop) keeps a
# bare `---` thematic break from being mistaken for a table.
_TABLE_SEPARATOR_RE = re.compile(r"^\|?[\s:|-]*-[\s:|-]*\|?$")
# Split on pipes that are not backslash-escaped — primer/docx_to_md.py
# escapes a literal pipe inside a cell as `\|`, so this is the matching
# unescape.
_UNESCAPED_PIPE_RE = re.compile(r"(?<!\\)\|")

# JLBC's memo has exactly ONE heading level plus bold run-in labels
# (`Policy Issues – ...`, `BUDS Table: ...`). `#`/`##` map to the section
# heading; anything deeper maps to the run-in label, which is what the
# third level actually is in this house style (spec M8).
_SECTION_HEADING_DEPTH = 2


def _is_table_row(line: str) -> bool:
    return "|" in line


def _split_row(line: str) -> list[str]:
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|") and not inner.endswith("\\|"):
        inner = inner[:-1]
    return [c.strip().replace("\\|", "|") for c in _UNESCAPED_PIPE_RE.split(inner)]


def _add_runs(paragraph, text: str) -> None:
    """Write `text` into a paragraph, turning **…** into bold runs.

    Everything outside the markers is emitted as-is, so an unmatched `**`
    stays visible rather than eating the rest of the line.
    """
    position = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _add_table(doc: DocumentT, rows: list[list[str]]) -> None:
    """Render collected pipe-table rows as a real Word table.

    Short rows are PADDED rather than dropped, matching what
    primer/docx_to_md.py does in the other direction: a ragged row is a
    formatting slip in the model's output, not a reason to lose the cell
    values it does contain.
    """
    columns = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(columns):
            value = row[column] if column < len(row) else ""
            paragraph = cells[column].paragraphs[0]
            _add_runs(paragraph, value)
            if index == 0:
                # Header row: bold every run, including ones **…** already
                # bolded (idempotent).
                for run in paragraph.runs:
                    run.bold = True


def _add_heading(doc: DocumentT, level: int, text: str) -> None:
    """A section heading, or — below the memo's one heading tier — the bold
    run-in label that is really its third level."""
    if level <= _SECTION_HEADING_DEPTH:
        paragraph = doc.add_paragraph(style="Header")
    else:
        paragraph = doc.add_paragraph()
    _add_runs(paragraph, text)
    # BOLD ON THE RUNS, NEVER ON THE STYLE. The memo block's DATE/TO/FROM/
    # SUBJECT labels share the `Header` paragraph style and are NOT bold in
    # the reference; putting bold on the style bolds them as a side effect.
    for run in paragraph.runs:
        run.bold = True


def render_body(doc: DocumentT, body_markdown: str) -> None:
    lines = body_markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        # Classify headings and bullets FIRST, because their text may
        # legitimately contain a pipe: "- Agency | Amount" is a bullet, and
        # treating it as a table header row (which it structurally
        # resembles, if the next line happens to be dashes) produced a
        # malformed table whose first cell was the literal "- Agency" —
        # the list item gone, the marker showing as garbled text. No text
        # technically vanished, but a memo where a bullet became a broken
        # table is the "no silent drops" rule failing in spirit.
        heading = _HEADING_RE.match(stripped)
        bullet = None if heading else _BULLET_RE.match(stripped)

        # A table row is only a table row when the NEXT line is a
        # separator; otherwise it is ordinary text containing pipe
        # characters and must survive as such.
        if (
            not heading
            and not bullet
            and _is_table_row(stripped)
            and index + 1 < len(lines)
            and _TABLE_SEPARATOR_RE.match(lines[index + 1].strip())
            and "|" in lines[index + 1]
        ):
            rows = [_split_row(stripped)]
            index += 2  # header + separator
            while index < len(lines) and _is_table_row(lines[index].strip()):
                rows.append(_split_row(lines[index].strip()))
                index += 1
            _add_table(doc, rows)
            continue

        if heading:
            _add_heading(doc, len(heading.group(1)), heading.group(2))
            index += 1
            continue

        if bullet:
            # `List Bullet` rather than the reference's `List Paragraph`:
            # it carries the bullet glyph through numbering.xml in
            # python-docx's default template, where `List Paragraph` does
            # not. The reference's bullets come from a numbering definition
            # its own file ships. Same rendered result, no hand-authored
            # numbering part.
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = BULLET_INDENT
            _add_runs(paragraph, bullet.group(1))
            index += 1
            continue

        # Everything else — blockquotes, numbered lists, code fences,
        # links, tables missing a separator — lands here VERBATIM. The
        # markup is visible but no content is lost.
        _add_runs(doc.add_paragraph(), line)
        index += 1
