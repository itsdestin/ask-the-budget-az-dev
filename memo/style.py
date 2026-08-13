"""Every measured value of the JLBC memo, and the chrome that surrounds
the body.

Measured from `samples/raw-docx/jlbc-staff-memorandum-style-reference.docx`
— the real FY 2027 Appropriations Report Round 1 instructions memo — not
eyeballed. `tests/test_jlbc_memo.py` reads most of these back out of that
file, so changing the house style is a fixture swap.

This module writes nothing and reads nothing. It builds a Document object
and hands it back. That is what lets `harness/documents.py` import it
without widening Invariant 7's blast radius (spec M1).
"""
from __future__ import annotations

from datetime import date as _date

from docx import Document
from docx.document import Document as DocumentT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

# --- copy, verbatim -------------------------------------------------------
MASTHEAD_TITLE = "Joint Legislative Budget Committee"
# NOT "Staff Memorandum" (spec M2). A Staff Memorandum is a specific JLBC
# work product with specific authorship; a machine-drafted document must
# not claim to be one, even though it carries the letterhead.
SUBTITLE = "Research Memorandum"
ADDRESS_LINES = (
    ("1716 West Adams", "Telephone: (602) 926-5491"),
    ("Phoenix, Arizona 85007", "azjlbc.gov"),
)
FOOTER_NOTE = "Generated with JLBC Agentic Search"
# The tool's name is stated ONCE and the suffix is derived from it, because
# the FROM row needs the bare name when there is no analyst to attribute.
# Deriving beats trimming the suffix at the use site: `lstrip` takes a SET
# of characters, so `SENDER_SUFFIX.lstrip(", ")` reads as "drop the comma
# and space" and actually yields "via JLBC Agentic Search" — a FROM row
# beginning "via". Measured in the interpreter, not assumed.
TOOL_NAME = "JLBC Agentic Search"
SENDER_SUFFIX = f", via {TOOL_NAME}"
# A visible placeholder, not an empty cell: an empty cell beside a label
# reads as a rendering bug (spec M4).
NO_RECIPIENT = "[Recipient(s)]"

# --- measured geometry ----------------------------------------------------
TOP_MARGIN = Inches(0.7)
BOTTOM_MARGIN = Inches(0.5)
SIDE_MARGIN = Inches(1.0)
HEADER_DISTANCE = Inches(0.75)
FOOTER_DISTANCE = Inches(0.6)

MASTHEAD_PT = Pt(14)
ADDRESS_PT = Pt(10)
BODY_PT = Pt(10.5)
FOOTER_PT = Pt(9)

TAB_STOP_TWIPS = Twips(7020)
LABEL_COL_TWIPS = Twips(1458)
VALUE_COL_TWIPS = Twips(7740)
BULLET_INDENT = Inches(0.1875)

# Word expresses border weight in EIGHTHS of a point, so 2.25pt is 18.
RULE_SIZE_EIGHTHS = "18"

BODY_FONT = "Calibri"


def today_long() -> str:
    """`August 12, 2026`. Built without `%-d`/`%#d`, which differ between
    glibc and the Windows CRT — this app ships to both."""
    now = _date.today()
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def new_document() -> DocumentT:
    """A blank document with the page set up and the styles corrected."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.left_margin = SIDE_MARGIN
    section.right_margin = SIDE_MARGIN
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE
    # The reference suppresses the page number on page 1, so the first
    # page needs its own header part.
    section.different_first_page_header_footer = True

    # Size on the STYLE, not per-run (spec M9): bullets, table cells and
    # any paragraph a later edit adds all inherit, instead of depending on
    # someone remembering to size each run. The reference sets 10.5 on
    # every run over a 12pt Normal; identical rendered result.
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_PT

    _fix_title_style(doc)
    _add_page_number_header(section)
    _add_footer_note(section)
    return doc


def _fix_title_style(doc: DocumentT) -> None:
    """python-docx's stock `Title` style is not JLBC's.

    It ships a blue bottom border (`w:pBdr`, sz 8, themeColor accent1) and
    `spacing after 300`. The reference has neither — verified by diffing
    both style definitions during planning. Left alone, the masthead draws
    a stray blue line immediately above the real 2.25pt rule, so the memo
    appears to have two.
    """
    title = doc.styles["Title"]
    title.font.size = MASTHEAD_PT
    title.font.bold = True
    title.font.name = BODY_FONT
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    pPr = title.element.get_or_add_pPr()
    for border in pPr.findall(qn("w:pBdr")):
        pPr.remove(border)


def _add_page_number_header(section) -> None:
    """`- 5 -` on page 2 onward; nothing on page 1."""
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("- ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    paragraph.add_run(" -")


def _add_footer_note(section) -> None:
    """The disclosure (spec M3), on every page.

    In the FOOTER rather than the body because a body line sits in the
    analyst's prose and gets deleted; a footer travels with the document,
    survives printing, and costs no vertical space. `different_first_page`
    is set for the header's sake, which means page 1 has its own footer
    part and needs the note written into it separately.
    """
    for footer in (section.footer, section.first_page_footer):
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(FOOTER_NOTE)
        run.font.size = FOOTER_PT


def add_masthead(doc: DocumentT, *, subtitle: str = SUBTITLE) -> None:
    doc.add_paragraph(MASTHEAD_TITLE, style="Title")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run(subtitle)
    run.bold = True
    run.font.size = MASTHEAD_PT

    for left, right in ADDRESS_LINES:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            TAB_STOP_TWIPS, WD_TAB_ALIGNMENT.LEFT
        )
        run = paragraph.add_run(f"{left}\t{right}")
        run.font.size = ADDRESS_PT


def add_rule(doc: DocumentT) -> None:
    """The line under the letterhead.

    The reference draws it as a VML `pict` at 2.25pt stroke. A paragraph
    bottom border is visually identical, is plain WordprocessingML, and
    needs no embedded drawing part (spec M10).
    """
    paragraph = doc.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), RULE_SIZE_EIGHTHS)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    pPr.append(borders)


MEMO_LABELS = ("DATE:", "TO:", "FROM:", "SUBJECT:")


def add_memo_block(
    doc: DocumentT,
    *,
    date: str,
    recipient: str,
    sender: str,
    subject: str,
) -> None:
    """The DATE / TO / FROM / SUBJECT block.

    A borderless two-column table with a blank spacer row between each
    pair, exactly as the reference builds it. It is invisible to
    `Document.paragraphs`, which is why a first pass at measuring this
    document missed it entirely.
    """
    # No sender degrades to the tool's name alone, never a dangling comma:
    # an unnameable analyst should lose attribution, not receive a
    # malformed memo.
    sender_line = f"{sender}{SENDER_SUFFIX}" if sender else TOOL_NAME
    values = (date, recipient or NO_RECIPIENT, sender_line, subject)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    # `autofit = False` alone does not hold the widths, and the reference
    # carries an explicit fixed layout, so this does too. It is what tells
    # WORD to honour the widths at render time — a thing no python-docx
    # getter can observe, so no test here can pin it.
    #
    # MEASURED on python-docx 1.2.0, correcting the plan: the 2743200-EMU
    # failure it attributes to a missing `tblLayout` does not come from
    # that. Omitting this element changes neither `columns[].width` nor
    # `cells[].width` on a save/reload. 2743200 is a CELL width, and what
    # produces it is creating the rows UP FRONT (`add_table(rows=7, …)`)
    # and then setting only the columns: the cells keep Word's even split,
    # the labels wrap, and the column widths still read correctly, so the
    # obvious assertion misses it. Hence `rows=0` below, with each row
    # added afterwards — that alone is what sizes the cells.
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    table.columns[0].width = LABEL_COL_TWIPS
    table.columns[1].width = VALUE_COL_TWIPS

    # The per-cell widths below are redundant on this construction — the
    # cells inherit the gridCol widths already, verified by deleting these
    # lines and watching every test stay green. They are kept because they
    # state the intent at the point a future edit is most likely to change
    # how the rows are built, which is the change that does break it.
    for index, (label, value) in enumerate(zip(MEMO_LABELS, values)):
        if index:
            spacer = table.add_row()
            spacer.cells[0].width = LABEL_COL_TWIPS
            spacer.cells[1].width = VALUE_COL_TWIPS
        row = table.add_row()
        row.cells[0].width = LABEL_COL_TWIPS
        row.cells[1].width = VALUE_COL_TWIPS
        # `Header` style on the label cell, matching the reference. Bold
        # is deliberately NOT applied — see the test that pins this.
        label_paragraph = row.cells[0].paragraphs[0]
        label_paragraph.style = doc.styles["Header"]
        label_paragraph.add_run(label)
        row.cells[1].paragraphs[0].add_run(value)
