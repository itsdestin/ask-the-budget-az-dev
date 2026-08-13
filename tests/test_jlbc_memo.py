"""The renderer is measured against the committed reference memo, not
against remembered numbers.

`REFERENCE` is `samples/raw-docx/jlbc-staff-memorandum-style-reference.docx`
— the real FY 2027 Appropriations Report Round 1 instructions memo. Where
a value can be read out of it, these tests read it, so a future JLBC style
change is a fixture swap rather than a code rewrite.

Two structural elements of that document are INVISIBLE to
`Document.paragraphs` and were missed on a first pass: the horizontal rule
is a VML `pict`, and the DATE/TO/FROM/SUBJECT block is a borderless table.
Anything that needs to see them must walk `document.body` elements.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import memo
from memo import style

REFERENCE = (
    Path(__file__).resolve().parents[1]
    / "samples"
    / "raw-docx"
    / "jlbc-staff-memorandum-style-reference.docx"
)


@pytest.fixture(scope="module")
def reference():
    return Document(str(REFERENCE))


@pytest.fixture
def rendered():
    return memo.render(
        "Body text.",
        subject="FY 2027 AHCCCS Appropriations Summary",
        sender="Destin Jarrett",
        recipient="",
        date="August 12, 2026",
    )


def test_page_margins_match_the_reference_to_the_emu(rendered, reference):
    ours, theirs = rendered.sections[0], reference.sections[0]
    assert ours.top_margin == theirs.top_margin
    assert ours.bottom_margin == theirs.bottom_margin
    assert ours.left_margin == theirs.left_margin
    assert ours.right_margin == theirs.right_margin
    assert ours.header_distance == theirs.header_distance
    assert ours.footer_distance == theirs.footer_distance


def test_the_masthead_carries_the_letterhead_with_the_research_subtitle(rendered):
    texts = [p.text for p in rendered.paragraphs[:4]]
    assert texts[0] == "Joint Legislative Budget Committee"
    assert texts[1] == "Research Memorandum"
    assert texts[2].startswith("1716 West Adams\t")
    assert texts[3].startswith("Phoenix, Arizona 85007\t")


def test_the_masthead_lines_are_14pt_bold_centered(rendered):
    title_style = rendered.styles["Title"]
    assert title_style.font.size == Pt(14)
    assert title_style.font.bold is True
    assert title_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    subtitle = rendered.paragraphs[1]
    assert subtitle.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert subtitle.runs[0].bold is True
    assert subtitle.runs[0].font.size == Pt(14)


def test_the_title_styles_default_blue_border_and_spacing_are_stripped(rendered):
    """python-docx's default `Title` style ships a blue bottom border
    (sz 8, accent1) and `spacing after 300`. The reference has neither, and
    leaving them draws a stray blue line directly above the real 2.25pt
    rule — two lines where the memo has one. Verified against a stock
    `Document()` during planning."""
    assert "pBdr" not in rendered.styles["Title"].element.xml
    assert rendered.styles["Title"].paragraph_format.space_after == Pt(0)


def test_the_address_lines_use_a_left_tab_stop_at_7020_twips(rendered):
    for index in (2, 3):
        stops = list(rendered.paragraphs[index].paragraph_format.tab_stops)
        assert len(stops) == 1
        assert stops[0].position == style.TAB_STOP_TWIPS
        assert rendered.paragraphs[index].runs[0].font.size == Pt(10)


def test_the_rule_is_a_225pt_paragraph_bottom_border(rendered):
    rule = rendered.paragraphs[4]
    assert "pBdr" in rule._p.xml
    assert 'w:sz="18"' in rule._p.xml  # 18 eighths of a point = 2.25pt


def test_the_footer_note_appears_on_the_first_page_too(rendered):
    section = rendered.sections[0]
    assert section.different_first_page_header_footer is True
    for footer in (section.footer, section.first_page_footer):
        assert footer.paragraphs[0].text == "Generated with JLBC Agentic Search"
        assert footer.paragraphs[0].runs[0].font.size == Pt(9)


def test_later_pages_carry_a_page_number_and_the_first_page_does_not(rendered):
    section = rendered.sections[0]
    assert "PAGE" in section.header.paragraphs[0]._p.xml
    assert section.first_page_header.paragraphs[0].text.strip() == ""


def test_body_text_is_calibri_105pt_on_the_normal_style(rendered):
    """Set on the style, not per-run, so bullets and table cells inherit
    (spec M9). The reference sets 10.5 on each run over a 12pt Normal;
    same rendered result, one place to change."""
    normal = rendered.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size == Pt(10.5)


def test_the_memo_block_is_a_borderless_7x2_table(rendered, reference):
    table = rendered.tables[0]
    theirs = reference.tables[0]
    assert len(table.rows) == len(theirs.rows) == 7
    assert len(table.columns) == 2
    assert [c.width for c in table.columns] == [c.width for c in theirs.columns]
    assert "tblBorders" not in table.style.element.xml


def test_every_memo_block_CELL_carries_its_own_width(rendered, reference):
    """Column widths alone do not protect the thing that actually wraps.

    Measured on python-docx 1.2.0, and it corrects the plan's account of
    this trap. Build the table with its rows already present
    (`add_table(rows=7, cols=2)`) and set only the columns, and every
    CELL round-trips at 2743200 EMU — Word's even split — while
    `columns[].width` still reads the intended 925830 / 4914900. The
    labels wrap and `test_the_memo_block_is_a_borderless_7x2_table`
    stays green throughout. This test is what goes red.

    It asserts the rendered property, not the mechanism: with the shipped
    `rows=0` + `add_row()` construction the cells inherit the gridCol
    widths on their own, so deleting the explicit `cells[…].width` lines
    in `add_memo_block` does NOT fail this — verified by deleting them.
    What fails it is a refactor to rows-up-front.
    """
    ours = rendered.tables[0]
    theirs = reference.tables[0]
    expected = [c.width for c in theirs.rows[0].cells]
    assert expected == [style.LABEL_COL_TWIPS, style.VALUE_COL_TWIPS]
    for index, row in enumerate(ours.rows):
        assert [c.width for c in row.cells] == expected, f"row {index} unsized"


def test_the_memo_block_labels_are_in_order_and_not_bold(rendered):
    """The labels share the `Header` paragraph style with section
    headings, and in the reference they are NOT bold — only the headings
    are, via direct run formatting. A future edit that puts bold on the
    style instead of the run bolds these as a side effect (spec M8)."""
    table = rendered.tables[0]
    labels = [table.rows[i].cells[0].text for i in (0, 2, 4, 6)]
    assert labels == ["DATE:", "TO:", "FROM:", "SUBJECT:"]
    for row in (0, 2, 4, 6):
        paragraph = table.rows[row].cells[0].paragraphs[0]
        assert paragraph.style.name == "Header"
        assert paragraph.runs[0].bold is not True


def test_the_subject_row_carries_the_documents_title(rendered):
    assert (
        rendered.tables[0].rows[6].cells[1].text
        == "FY 2027 AHCCCS Appropriations Summary"
    )


def test_there_is_no_separate_title_line_in_the_body(rendered):
    """The reference has none, and Word's `Title` style is already spent
    on the masthead. The subject IS the title (spec M4)."""
    body_texts = [p.text for p in rendered.paragraphs]
    assert body_texts.count("FY 2027 AHCCCS Appropriations Summary") == 0


def test_an_absent_recipient_renders_a_visible_placeholder(rendered):
    assert rendered.tables[0].rows[2].cells[1].text == "[Recipient(s)]"


def test_a_supplied_recipient_is_used_verbatim():
    doc = memo.render(
        "Body.", subject="S", sender="A. Analyst", recipient="Director Smith"
    )
    assert doc.tables[0].rows[2].cells[1].text == "Director Smith"


def test_the_sender_row_names_the_analyst_and_the_tool(rendered):
    assert (
        rendered.tables[0].rows[4].cells[1].text
        == "Destin Jarrett, via JLBC Agentic Search"
    )


def test_an_unknown_analyst_leaves_the_tool_alone_on_the_from_line():
    """Degrades to the tool's name rather than a dangling comma. An
    unnameable user should lose attribution, not get a malformed memo —
    the same posture `app.identity.current_user` already takes."""
    doc = memo.render("Body.", subject="S", sender="")
    assert doc.tables[0].rows[4].cells[1].text == "JLBC Agentic Search"


def test_the_date_defaults_to_today_in_long_form():
    from memo.style import today_long

    doc = memo.render("Body.", subject="S", sender="A")
    assert doc.tables[0].rows[0].cells[1].text == today_long()


def _body_paragraphs(doc):
    """Paragraphs after the chrome. The masthead is 4 paragraphs, then the
    rule, then the memo block (a table, which does not appear here).

    VERIFIED against what Tasks 1-2 actually emit rather than taken from
    the plan on trust: `render()` with an empty body produces exactly five
    paragraphs — Title, subtitle, two address lines, rule — and the memo
    block contributes none, because `Document.paragraphs` does not descend
    into table cells. So the body starts at index 5.
    """
    return doc.paragraphs[5:]


def test_a_heading_becomes_a_bold_header_styled_paragraph():
    doc = memo.render("## Policy Issues\n\nText.", subject="S", sender="A")
    heading = [p for p in _body_paragraphs(doc) if p.text == "Policy Issues"][0]
    assert heading.style.name == "Header"
    assert heading.runs[0].bold is True


def test_bold_is_never_put_on_the_header_style_itself():
    """The memo block's labels share this style and must stay unbold."""
    doc = memo.render("## Policy Issues", subject="S", sender="A")
    assert doc.styles["Header"].font.bold is not True


def test_a_deep_heading_becomes_a_bold_run_in_label():
    """`###` and deeper map to the memo's third level, which is a bold
    run-in label (`BUDS Table: ...`), not another heading tier."""
    doc = memo.render("### BUDS Table", subject="S", sender="A")
    label = [p for p in _body_paragraphs(doc) if p.text == "BUDS Table"][0]
    assert label.style.name == "Normal"
    assert label.runs[0].bold is True


def test_a_bullet_is_indented_to_the_house_measure():
    doc = memo.render("- One item", subject="S", sender="A")
    bullet = [p for p in _body_paragraphs(doc) if p.text == "One item"][0]
    assert bullet.style.name == "List Bullet"
    assert bullet.paragraph_format.left_indent == style.BULLET_INDENT


def test_bold_markers_become_bold_runs():
    doc = memo.render("The **budget includes** $5.", subject="S", sender="A")
    paragraph = [p for p in _body_paragraphs(doc) if p.text.startswith("The ")][0]
    assert [(r.text, r.bold) for r in paragraph.runs] == [
        ("The ", None),
        ("budget includes", True),
        (" $5.", None),
    ]


def test_a_pipe_table_becomes_a_real_table_after_the_memo_block():
    doc = memo.render(
        "| Agency | Amount |\n|---|---|\n| ADC | $5 |", subject="S", sender="A"
    )
    assert len(doc.tables) == 2  # memo block + this one
    body_table = doc.tables[1]
    assert body_table.style.name == "Table Grid"
    assert [c.text for c in body_table.rows[0].cells] == ["Agency", "Amount"]
    assert body_table.rows[0].cells[0].paragraphs[0].runs[0].bold is True


def test_a_bullet_containing_a_pipe_stays_a_bullet():
    """Regression, carried over from harness/documents.py: classifying
    table rows before bullets turned `- Agency | Amount` into a malformed
    table whose first cell read `- Agency`."""
    doc = memo.render("- Agency | Amount\n|---|---|", subject="S", sender="A")
    assert len(doc.tables) == 1  # the memo block only
    bullet = [p for p in _body_paragraphs(doc) if "Agency" in p.text][0]
    assert bullet.style.name == "List Bullet"


def test_unrecognized_markup_survives_verbatim():
    """No silent drops. An analyst who receives a memo with a section
    quietly missing has no way to know it happened; a blockquote showing
    its `>` is a far better failure."""
    doc = memo.render("> quoted line\n\n1. numbered", subject="S", sender="A")
    texts = [p.text for p in _body_paragraphs(doc)]
    assert "> quoted line" in texts
    assert "1. numbered" in texts


def test_memo_package_imports_are_allowlisted():
    """The transitive half of Invariant 7 (spec M1). `harness/documents.py`
    is allowed to import `memo`; that concession is only safe while `memo`
    itself cannot reach the shared drive. Adding an entry here has to be a
    conscious edit, exactly like the allowlist it backs."""
    import ast

    allowed = {"__future__", "dataclasses", "datetime", "re", "typing", "docx", "memo"}
    package = Path(memo.__file__).parent
    roots: set[str] = set()
    for source in package.glob("*.py"):
        tree = ast.parse(source.read_text(encoding="utf-8"))
        for node in ast.walk(tree):
            if isinstance(node, ast.Import):
                roots.update(a.name.split(".")[0] for a in node.names)
            elif isinstance(node, ast.ImportFrom) and node.module and node.level == 0:
                roots.add(node.module.split(".")[0])
    assert roots <= allowed, f"unexpected imports: {sorted(roots - allowed)}"


def test_materialize_renders_a_memo_and_keeps_the_title_in_properties(
    tmp_path, monkeypatch
):
    from harness import documents

    monkeypatch.setenv(documents.DOCUMENTS_DIR_ENV, str(tmp_path))
    documents.reset_registry()
    _token, path = documents.materialize(
        "FY 2027 Summary",
        "## Policy Issues\n\nText.",
        "docx",
        user="djarrett",
        sender="Destin Jarrett",
        recipient="Director Smith",
    )
    doc = Document(str(path))
    assert doc.core_properties.title == "FY 2027 Summary"
    assert doc.paragraphs[0].text == "Joint Legislative Budget Committee"
    assert doc.tables[0].rows[6].cells[1].text == "FY 2027 Summary"
    assert doc.tables[0].rows[2].cells[1].text == "Director Smith"
    assert (
        doc.tables[0].rows[4].cells[1].text
        == "Destin Jarrett, via JLBC Agentic Search"
    )


def test_the_markdown_format_path_is_untouched(tmp_path, monkeypatch):
    """`format="md"` stays byte-faithful to what the model wrote (spec
    M12). It is the escape hatch for an analyst who wants the text without
    the formatting, and round-tripping it could lose a construct."""
    from harness import documents

    monkeypatch.setenv(documents.DOCUMENTS_DIR_ENV, str(tmp_path))
    documents.reset_registry()
    body = "> quoted\n\n1. numbered\n"
    _token, path = documents.materialize("T", body, "md")
    assert path.read_text(encoding="utf-8") == body
