from pathlib import Path

import fitz

from ingest.inspection import inspect_source


def _pdf(tmp_path: Path, *, pages: int, text: bool) -> Path:
    doc = fitz.open()
    for _ in range(pages):
        page = doc.new_page()
        if text:
            page.insert_text((72, 72), "The Baseline includes $1,000,000 for X.")
    path = tmp_path / "f.pdf"
    doc.save(path)
    doc.close()
    return path


def test_reports_format_and_page_count(tmp_path):
    got = inspect_source(_pdf(tmp_path, pages=3, text=True))
    assert got.source_format == "pdf"
    assert got.pages == 3


def test_detects_a_text_layer(tmp_path):
    assert inspect_source(_pdf(tmp_path, pages=1, text=True)).has_text_layer is True


def test_detects_the_absence_of_a_text_layer(tmp_path):
    assert inspect_source(_pdf(tmp_path, pages=1, text=False)).has_text_layer is False


def test_inspection_reports_nothing_about_tagging(tmp_path):
    """A regression guard, not a capability test.

    Tagging was measured across every OpenDataLoader-first document and does
    not predict extraction success: an UNTAGGED 639-page Executive Budget
    scores 92.2% through OpenDataLoader, while the one document that fails IS
    tagged. A future edit that re-adds this field will re-add the rule that
    consumes it, which diverts a healthy document to a slower extractor.
    """
    got = inspect_source(_pdf(tmp_path, pages=1, text=True))
    assert not hasattr(got, "has_structure_tree")


def test_docx_has_no_page_count(tmp_path):
    """DOCX has no pages at rest. None, not 0 -- 0 would read as "empty"."""
    import docx

    d = docx.Document()
    d.add_paragraph("Section 1. Appropriations.")
    path = tmp_path / "bill.docx"
    d.save(path)

    got = inspect_source(path)
    assert got.source_format == "docx"
    assert got.pages is None
    assert got.has_text_layer is True


def test_an_unreadable_file_does_not_raise(tmp_path):
    """A truncated download must not take the worker thread down.

    pdfium rejects shapes PyMuPDF tolerates and vice versa; whatever the
    reason, an inspection failure means "we learned nothing", which is a
    valid answer that starts the ladder at rung 1.

    `has_text_layer` must be None, not False, on this path -- False means
    we POSITIVELY determined there is no text layer (a scan). Reporting
    False here instead would be indistinguishable from `import fitz`
    failing outright (e.g. a broken Windows bundle), which would silently
    route every PDF in the corpus to OCR-only. See ingest/ladder.py.
    """
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4\nnot really a pdf")

    got = inspect_source(path)
    assert got.pages is None
    assert got.has_text_layer is None


def test_an_unreadable_docx_does_not_raise(tmp_path):
    """Budget bills in this corpus are DOCX-only by design (spec T5), so a
    truncated DOCX download is a realistic production case, not a curiosity.

    `docx.Document()` raises `docx.opc.exceptions.PackageNotFoundError` on
    garbage bytes -- a distinct failure mode from the PDF branch above, and
    one the existing "unreadable file" test never exercises because it only
    ever builds a `.pdf` path.
    """
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a zip, not a docx, just garbage")

    got = inspect_source(path)
    assert got.source_format == "docx"
    assert got.pages is None
    assert got.has_text_layer is None


def test_a_scanned_pdf_with_only_whitespace_glyphs_has_no_text_layer(tmp_path):
    """Pins the `.strip()` in the has_text check, not just its absence.

    The existing "no text layer" fixture inserts nothing at all, so
    `get_text()` already returns "" and the strip() is a no-op there. A real
    scanned page often carries a few stray whitespace glyphs -- this fixture
    reproduces that shape directly, so a version of the code that dropped
    `.strip()` and treated any non-empty string as "has text" would flip this
    to True.
    """
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "   \n\t  ")
    path = tmp_path / "scanned.pdf"
    doc.save(path)
    doc.close()

    got = inspect_source(path)
    assert got.has_text_layer is False


def test_detects_a_text_layer_that_lives_only_in_table_cells(tmp_path):
    """`has_text` also scans table cells (spec T5), separately from
    paragraphs -- a DOCX budget bill can carry its substantive text in a
    table (e.g. an appropriations line-item grid) with no body paragraphs at
    all. The existing DOCX fixture only ever puts text in a paragraph, so it
    cannot catch a version of the code that dropped the table-cell branch.
    """
    import docx

    d = docx.Document()
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "HB 2001"
    table.rows[0].cells[1].text = "$1,000,000"
    path = tmp_path / "table-only.docx"
    d.save(path)

    got = inspect_source(path)
    assert got.has_text_layer is True
