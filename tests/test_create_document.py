"""harness/documents.py + app/routes/documents.py (Plan 4 Task 4).

`create_document` is the model's ONLY write primitive, so these tests are
mostly about what it is NOT allowed to do:

1. **Invariant 7** — nothing lands under the shared data directory. One
   process serves a whole office off a network share; a confused or
   prompt-injected model that could write there would corrupt state for
   everyone. Guarded behaviorally (the share tmp dir stays empty after a
   materialize) AND structurally (the module cannot import the module
   that knows where the share is).
2. **No model-supplied string reaches a filesystem path.** The model
   supplies a title; the sanitizer decides the filename and a
   containment check re-verifies the result. The download token is a
   dict key, never a path component, so a token shaped like
   `../../settings.json` is simply an unknown key.
3. **No silent drops.** An analyst who asks for a memo must not receive
   one that quietly lost a section, so every unrecognized construct has
   to survive as visible text.

Round-trips go through python-docx on the READ side too — asserting on
the written file rather than trusting the writer's own bookkeeping.

Route tests live in this file rather than a sibling because every one of
them needs an artifact materialized through the same in-process
registry; splitting them would duplicate the whole fixture.
"""
from __future__ import annotations

import ast
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import create_app
from app.search_provider import StubSearchProvider
from harness import documents
from store.config import data_dir

MODULE_SOURCE_PATH = Path(__file__).resolve().parent.parent / "harness" / "documents.py"

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@pytest.fixture(autouse=True)
def isolated_storage(tmp_path, monkeypatch):
    """Every test gets its own documents dir, its own (empty) share, and a
    clean registry.

    JLBC_DATA_DIR is repointed even though nothing here should touch it —
    that is exactly the point: if this module ever starts writing to the
    share, the Invariant 7 test below sees it in tmp_path instead of
    silently scribbling in the developer's real data directory.
    """
    monkeypatch.setenv(documents.DOCUMENTS_DIR_ENV, str(tmp_path / "local-docs"))
    monkeypatch.setenv("JLBC_DATA_DIR", str(tmp_path / "share"))
    documents.reset_registry()
    yield
    documents.reset_registry()


def client() -> TestClient:
    # Stub provider so constructing the app never probes for a real corpus;
    # these tests are about the documents route only.
    return TestClient(create_app(provider=StubSearchProvider(), static_dir=None))


def _all_text(doc: Document) -> list[str]:
    """Every visible string in the document — body paragraphs AND table
    cells. Table text is NOT in doc.paragraphs, so a no-silent-drops check
    that only walked paragraphs would miss content that moved into a
    table."""
    out = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            out.extend(cell.text for cell in row.cells)
    return out


# The .docx is a JLBC memo now (see memo/), so the model's body no longer
# starts at paragraph 0. `render()` emits five chrome paragraphs first —
# the Title-styled masthead, the "Research Memorandum" subtitle, two
# address lines, and the paragraph carrying the horizontal rule — followed
# by the DATE/TO/FROM/SUBJECT block, which is a TABLE and therefore
# contributes no paragraphs at all.
#
# VERIFIED against what render() actually emits rather than counted off
# the source: `[p.text for p in memo.render("", subject="S").paragraphs]`
# is exactly those five.
_CHROME_PARAGRAPHS = 5


def _body_paragraphs(doc: Document) -> list:
    """Paragraphs of the model's body, ignoring chrome and blank spacers."""
    return [p for p in doc.paragraphs[_CHROME_PARAGRAPHS:] if p.text.strip()]


def _body_tables(doc: Document) -> list:
    """Tables from the model's body.

    `tables[0]` is always the memo block, so a body table that used to be
    `tables[0]` is now `tables[1]`. Going through this helper keeps that
    offset stated in ONE place — and keeps `len(...) == 1` meaning "the
    body produced one table", which is the property these tests are
    actually about.
    """
    return doc.tables[1:]


# ---------------------------------------------------------------------------
# materialize — round trip
# ---------------------------------------------------------------------------


def test_materialize_writes_a_docx_that_opens():
    token, path = documents.materialize(
        "FY2027 AHCCCS memo", "# Summary\n\nSpending rose.", "docx", user="analyst1"
    )
    assert token
    assert path.is_file()
    assert path.suffix == ".docx"
    assert documents.documents_dir() in path.parents
    assert "Spending rose." in _all_text(Document(str(path)))


def test_materialize_writes_markdown_byte_faithfully():
    body = "# Heading\n\n- one\n\n> quoted\n\nplain — em dash"
    _, path = documents.materialize("Notes", body, "md", user="analyst1")
    assert path.suffix == ".md"
    # Byte-faithful: the .md path is a passthrough, so what the model wrote
    # is exactly what the analyst downloads.
    assert path.read_text(encoding="utf-8") == body


def test_an_unknown_format_is_rejected():
    with pytest.raises(ValueError):
        documents.materialize("T", "body", "pdf", user="analyst1")


def test_lookup_returns_the_artifact_and_records_the_user():
    token, path = documents.materialize("T", "body", "md", user="analyst1")
    artifact = documents.lookup(token)
    assert artifact is not None
    assert artifact.path == path
    assert artifact.user == "analyst1"
    assert artifact.media_type.startswith("text/markdown")


def test_unknown_token_lookup_returns_none():
    assert documents.lookup("nope") is None


# ---------------------------------------------------------------------------
# Markdown subset — verified by reading the .docx back
# ---------------------------------------------------------------------------


def _render(body: str) -> Document:
    _, path = documents.materialize("Doc title", body, "docx", user="u")
    return Document(str(path))


def test_the_title_is_carried_on_the_page_by_the_memo_blocks_subject_row():
    """Same property as before — the title is IN the document, not only in
    its filename, so the artifact is still titled once it is printed or
    pasted into an email. Only WHERE moved: Word's `Title` style is spent
    on the JLBC masthead now, and the memo has no separate title line
    (spec M4), so the title rides the SUBJECT row of the memo block."""
    doc = _render("body text")
    assert doc.paragraphs[0].text == "Joint Legislative Budget Committee"
    assert doc.paragraphs[0].style.name == "Title"
    assert doc.tables[0].rows[6].cells[0].text == "SUBJECT:"
    assert doc.tables[0].rows[6].cells[1].text == "Doc title"
    # ...and it appears exactly once. A body copy would be a second title
    # competing with the subject row.
    assert [p.text for p in doc.paragraphs].count("Doc title") == 0


def test_hash_headings_become_the_memos_heading_tiers():
    """The stock `Heading 1/2/3` ladder is gone because the memo does not
    have one: JLBC's house style is a single `Header`-styled section
    heading plus a bold run-in label beneath it (spec M8). `#`/`##` map to
    the heading; anything deeper maps to the run-in label. Every tier is
    still visually distinct from body text, which is what the old
    assertion was protecting."""
    doc = _render("# Top\n\n## Second\n\n### Third")
    paragraphs = _body_paragraphs(doc)
    assert [(p.text, p.style.name) for p in paragraphs] == [
        ("Top", "Header"),
        ("Second", "Header"),
        ("Third", "Normal"),
    ]
    assert all(run.bold for p in paragraphs for run in p.runs)


def test_dash_and_star_bullets_become_list_paragraphs():
    doc = _render("- first\n- second\n* third")
    styled = [(p.text, p.style.name) for p in _body_paragraphs(doc)]
    assert styled == [
        ("first", "List Bullet"),
        ("second", "List Bullet"),
        ("third", "List Bullet"),
    ]


def test_bold_markers_become_bold_runs_and_the_markers_disappear():
    doc = _render("Spending was **$2.4 million** last year.")
    para = _body_paragraphs(doc)[0]
    assert para.text == "Spending was $2.4 million last year."
    bolded = [r.text for r in para.runs if r.bold]
    assert bolded == ["$2.4 million"]


def test_bold_works_inside_headings_and_bullets():
    """Unchanged property: the `**` markers are consumed and the
    emphasized text really is bold. The heading half is asserted
    separately now because the memo's one heading tier is bold END TO END
    — `**` inside it is consumed but adds no contrast — while in body text
    only the marked span is bold, exactly as before."""
    heading, bullet = _body_paragraphs(
        _render("## A **bold** heading\n\n- a **bold** bullet")
    )
    assert "**" not in heading.text and "**" not in bullet.text

    assert heading.text == "A bold heading"
    assert [r.text for r in heading.runs if r.bold] == ["A ", "bold", " heading"]

    assert bullet.text == "a bold bullet"
    assert [r.text for r in bullet.runs if r.bold] == ["bold"]


def test_a_pipe_table_becomes_a_real_word_table():
    doc = _render(
        "| Agency | FY2027 |\n"
        "| --- | --- |\n"
        "| ADC | $1,000 |\n"
        "| AHCCCS | $2,000 |"
    )
    assert len(_body_tables(doc)) == 1
    table = _body_tables(doc)[0]
    assert [c.text for c in table.rows[0].cells] == ["Agency", "FY2027"]
    assert [c.text for c in table.rows[1].cells] == ["ADC", "$1,000"]
    assert [c.text for c in table.rows[2].cells] == ["AHCCCS", "$2,000"]


def test_a_table_cell_with_an_escaped_pipe_keeps_the_pipe():
    # primer/docx_to_md.py escapes literal pipes as \| when it WRITES a
    # markdown table; this is the matching unescape on the way back.
    doc = _render("| a | b |\n| --- | --- |\n| x \\| y | z |")
    assert _body_tables(doc)[0].rows[1].cells[0].text == "x | y"


def test_a_short_row_is_padded_rather_than_dropped():
    doc = _render("| a | b | c |\n| --- | --- | --- |\n| only |")
    row = _body_tables(doc)[0].rows[1]
    assert [c.text for c in row.cells] == ["only", "", ""]


@pytest.mark.parametrize(
    "first_line, expected_text, expected_style",
    [
        ("- Agency | Amount", "Agency | Amount", "List Bullet"),
        ("* Agency | Amount", "Agency | Amount", "List Bullet"),
        # `Heading 2` / `Heading 1` in the generic renderer; the memo has a
        # single `Header`-styled heading tier instead (spec M8). The
        # property under test — the line stays a heading rather than
        # becoming a table cell — is unchanged.
        ("## Costs | FY2027", "Costs | FY2027", "Header"),
        ("# Revenue | Expense", "Revenue | Expense", "Header"),
    ],
)
def test_a_piped_bullet_or_heading_is_not_mistaken_for_a_table_header(
    first_line, expected_text, expected_style
):
    """A pipe inside bullet or heading TEXT used to misclassify the line as
    a table header row, so the bullet marker survived as a literal "- Agency"
    cell and the list item was gone. No text technically vanished, but the
    analyst got a malformed table where a bullet should be — a silent drop in
    spirit. Heading/bullet now win the classification."""
    doc = _render(f"{first_line}\n| --- |\n| ADC | 1 |")
    first = _body_paragraphs(doc)[0]
    assert (first.text, first.style.name) == (expected_text, expected_style)
    assert _body_tables(doc) == []
    # The stray table fragments still survive verbatim as plain text.
    rendered = [t.strip() for t in _all_text(doc)]
    assert "| --- |" in rendered and "| ADC | 1 |" in rendered


def test_a_genuine_table_after_a_heading_still_renders():
    # Regression guard for the fix above: the heading and the table are on
    # separate lines, which is the normal shape, and must be unaffected.
    doc = _render("## Costs\n\n| Agency | FY2027 |\n| --- | --- |\n| ADC | $1 |")
    assert _body_paragraphs(doc)[0].style.name == "Header"
    assert [c.text for c in _body_tables(doc)[0].rows[1].cells] == ["ADC", "$1"]


def test_text_after_a_table_returns_to_normal_paragraphs():
    doc = _render("| a |\n| --- |\n| x |\n\nAfter the table.")
    assert _body_tables(doc)[0].rows[1].cells[0].text == "x"
    assert _body_paragraphs(doc)[-1].text == "After the table."


# ---------------------------------------------------------------------------
# No silent drops
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "line",
    [
        "> a blockquote",
        "1. a numbered item",
        "```python",
        "    indented code",
        "---",
        "[a link](http://example.com)",
        "| a table row with no separator |",
        "# ",
    ],
)
def test_unrecognized_markdown_survives_as_visible_text(line):
    """The hard requirement: an analyst must never get a document that
    quietly lost a section. Anything the renderer does not understand is
    emitted verbatim as a plain paragraph."""
    doc = _render(f"before\n\n{line}\n\nafter")
    assert line.strip() in [t.strip() for t in _all_text(doc)]


def test_every_nonblank_line_of_a_mixed_document_appears_somewhere():
    body = (
        "# Heading one\n"
        "Intro paragraph.\n"
        "- bullet one\n"
        "1. numbered item\n"
        "> blockquote\n"
        "```\ncode line\n```\n"
        "| Col A | Col B |\n"
        "| --- | --- |\n"
        "| v1 | v2 |\n"
        "Closing paragraph."
    )
    rendered = " \n ".join(_all_text(_render(body)))
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped or set(stripped) <= set("|- "):
            continue  # the table separator row is structure, not content
        # Heading/bullet markers are consumed by design; their TEXT is not.
        payload = stripped.lstrip("#-*> ").strip()
        if payload.startswith("|"):
            payload = payload.strip("|").split("|")[0].strip()
        assert payload in rendered, f"lost: {line!r}"


# ---------------------------------------------------------------------------
# Invariant 7 — never the shared data directory
# ---------------------------------------------------------------------------


def test_materialize_writes_nothing_under_the_shared_data_dir():
    share = data_dir()
    before = sorted(p.name for p in share.rglob("*"))
    documents.materialize("Memo", "# Body\n\n- a\n\n| x |\n| --- |\n| y |", "docx", user="u")
    documents.materialize("Memo", "body", "md", user="u")
    assert sorted(p.name for p in share.rglob("*")) == before


def test_the_documents_dir_is_not_inside_the_shared_data_dir():
    share = data_dir().resolve()
    local = documents.documents_dir().resolve()
    assert share != local and share not in local.parents


def test_documents_module_cannot_reach_the_shared_data_dir():
    """Structural half of Invariant 7, mirroring the import allowlist in
    tests/test_harness_tools.py: this module must not import store.config
    (or anything from ingest), so there is no code path by which it could
    learn the share's location, let alone write there."""
    allowed = {
        "__future__",
        "dataclasses",
        "os",
        "pathlib",
        "re",
        "secrets",
        "threading",
        "time",
        "typing",
        "unicodedata",
        "docx",
        # `memo` renders Markdown into a Word document and does nothing
        # else. Safe to allow because tests/test_jlbc_memo.py pins ITS
        # imports the same way, so the guarantee stays structural and
        # becomes transitive rather than becoming a promise.
        "memo",
    }
    tree = ast.parse(MODULE_SOURCE_PATH.read_text(encoding="utf-8"))
    roots: set[str] = set()
    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            roots.update(alias.name.split(".")[0] for alias in node.names)
        elif isinstance(node, ast.ImportFrom) and node.module and node.level == 0:
            roots.add(node.module.split(".")[0])
    assert roots <= allowed, f"unexpected imports: {sorted(roots - allowed)}"


# ---------------------------------------------------------------------------
# Tokens + filenames — nothing the model types reaches a path
# ---------------------------------------------------------------------------


def test_tokens_are_long_unique_and_not_derived_from_the_title():
    tokens = {documents.materialize("Same title", "b", "md", user="u")[0] for _ in range(20)}
    assert len(tokens) == 20
    for token in tokens:
        assert len(token) >= 32
        assert "Same" not in token


def test_two_documents_with_the_same_title_do_not_overwrite_each_other():
    t1, p1 = documents.materialize("Memo", "first body", "md", user="u")
    t2, p2 = documents.materialize("Memo", "second body", "md", user="u")
    assert p1 != p2
    assert p1.read_text(encoding="utf-8") == "first body"
    assert p2.read_text(encoding="utf-8") == "second body"
    assert documents.lookup(t1).path == p1
    assert documents.lookup(t2).path == p2


@pytest.mark.parametrize(
    "title",
    [
        "../../settings.json",
        "..\\..\\settings",
        "/etc/passwd",
        "C:\\Windows\\system32\\evil",
        'quote" ; rm -rf',
        "line\r\nbreak",
        "..",
        "   ",
        "***",
        "CON",
        "x" * 300,
    ],
)
def test_a_hostile_title_produces_a_safe_contained_filename(title):
    _, path = documents.materialize(title, "body", "md", user="u")
    root = documents.documents_dir().resolve()
    resolved = path.resolve()
    assert root in resolved.parents, f"{resolved} escaped {root}"
    assert path.is_file()
    name = path.name
    # A title made only of dots/spaces/symbols must fall back to a real
    # stem, not produce a file called "..md" or "   .md".
    assert path.stem.strip(". ")
    assert not any(ch in name for ch in '/\\:*?"<>|\r\n')
    assert len(name) <= 128
    assert name.rstrip(". ") == name  # Windows rejects trailing dots/spaces


def test_a_reserved_windows_device_name_is_not_used_bare():
    _, path = documents.materialize("NUL", "body", "md", user="u")
    assert path.stem.upper() != "NUL"


@pytest.mark.parametrize(
    "title, expected_stem",
    [
        # The filter is str.isalnum(), which is Unicode-aware, so letters
        # from any script survive. Pinned because the module comment claims
        # exactly this.
        ("预算报告", "预算报告"),
        ("Réunion budgétaire", "Réunion-budgétaire"),
        # ...and anything NOT alphanumeric or in "-_." is replaced, even
        # when a URL or a filesystem would have tolerated it.
        ("unicode~tilde", "unicode-tilde"),
        ("memo (draft)", "memo-draft"),
        ("A&B, Inc.", "A-B-Inc"),
    ],
)
def test_the_filename_allowlist_is_unicode_alphanumerics_plus_dash_dot_underscore(
    title, expected_stem
):
    _, path = documents.materialize(title, "body", "md", user="u")
    assert path.stem == expected_stem


def test_a_unicode_filename_downloads_with_an_rfc_5987_header():
    """The other half of the Unicode decision: a non-ASCII filename is
    legal on disk and legal in the header, because Starlette falls back to
    the `filename*=utf-8''…` form rather than emitting raw bytes."""
    token, _ = documents.materialize("预算报告", "body", "md", user="u")
    r = client().get(f"/api/documents/{token}")
    assert r.status_code == 200
    assert "filename*=utf-8''" in r.headers["content-disposition"]


# ---------------------------------------------------------------------------
# A failed write must never leave a registered token
# ---------------------------------------------------------------------------


@pytest.mark.parametrize("fmt", ["docx", "md"])
def test_a_failed_write_registers_no_token(fmt, monkeypatch):
    """Registration happens strictly AFTER the bytes land. A refactor that
    minted the token first would hand the model a download link to a file
    that does not exist (or, worse, to a half-written .docx Word refuses to
    open) — and the tool result would say ok:true.

    Reaching into the private registry on purpose: the public `lookup()`
    needs a token, and the whole point is that no token was ever returned.
    """
    if fmt == "docx":
        monkeypatch.setattr(
            documents, "_render_docx", lambda *a, **k: (_ for _ in ()).throw(OSError("disk full"))
        )
    else:
        def boom(*a, **k):
            raise OSError("disk full")

        monkeypatch.setattr(Path, "write_text", boom)

    with pytest.raises(OSError):
        documents.materialize("Memo", "# Body", fmt, user="u")
    assert documents._registry == {}


# ---------------------------------------------------------------------------
# GET /api/documents/{token}
# ---------------------------------------------------------------------------


def test_download_serves_the_docx_with_the_right_content_type():
    token, path = documents.materialize("FY2027 memo", "# Body", "docx", user="u")
    r = client().get(f"/api/documents/{token}")
    assert r.status_code == 200
    assert r.headers["content-type"] == DOCX_MEDIA_TYPE
    disposition = r.headers["content-disposition"]
    assert disposition.startswith("attachment")
    assert path.name in disposition
    assert r.content == path.read_bytes()


def test_download_serves_markdown_as_text():
    token, _ = documents.materialize("Notes", "# Body", "md", user="u")
    r = client().get(f"/api/documents/{token}")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/markdown")
    assert r.text == "# Body"


def test_unknown_token_is_404():
    r = client().get("/api/documents/definitely-not-a-token")
    assert r.status_code == 404
    assert r.json()["detail"]


@pytest.mark.parametrize(
    "token",
    [
        # Percent-encoded on purpose. A literal "../../settings.json" is
        # collapsed by the HTTP client before the request is even sent, so
        # it never reaches this route and would test nothing about it.
        "..%2F..%2Fsettings.json",
        "..%5C..%5Csettings.json",
        "%2e%2e%2f%2e%2e%2fsettings.json",
        "C:%5CWindows%5Csystem32%5Cdrivers%5Cetc%5Chosts",
        "%00settings.json",
    ],
)
def test_a_traversal_shaped_token_cannot_escape(token, tmp_path):
    """The token is a registry KEY, never a path component, so traversal
    is not filtered — it is impossible. A secret planted next to the
    documents dir stays unreachable."""
    (tmp_path / "share").mkdir(parents=True, exist_ok=True)
    (tmp_path / "share" / "settings.json").write_text('{"api_key": "SECRET"}')
    r = client().get(f"/api/documents/{token}")
    assert r.status_code == 404
    assert "SECRET" not in r.text


def test_a_traversal_token_reaching_the_route_is_just_an_unknown_key():
    """The case above proves the SYSTEM is safe; this one proves the route
    itself is, by using a token that survives URL normalization as a single
    path segment and therefore really does reach download_document()."""
    r = client().get("/api/documents/..%5C..%5Csettings.json")
    assert r.status_code == 404
    assert "Unknown download link" in r.json()["detail"]
    assert documents.lookup("..\\..\\settings.json") is None


def test_a_download_can_be_repeated():
    """DELIBERATELY NOT single-use. Browsers issue HEAD/range requests and
    users click links twice; a one-shot token would turn an ordinary
    second click into a broken download with no way to regenerate the
    file short of re-asking the model."""
    token, _ = documents.materialize("Memo", "body", "md", user="u")
    c = client()
    assert c.get(f"/api/documents/{token}").status_code == 200
    assert c.get(f"/api/documents/{token}").status_code == 200


def test_a_registered_token_whose_file_vanished_is_404_not_a_crash():
    token, path = documents.materialize("Memo", "body", "md", user="u")
    path.unlink()
    r = client().get(f"/api/documents/{token}")
    assert r.status_code == 404


def test_the_documents_route_is_registered_before_the_spa_catch_all():
    # Regression guard: app/main.py's `/{path:path}` route swallows anything
    # registered after it, which would turn every download into index.html.
    paths = [getattr(r, "path", "") for r in create_app(
        provider=StubSearchProvider(), static_dir=None
    ).routes]
    assert "/api/documents/{token}" in paths
    assert paths.index("/api/documents/{token}") < paths.index("/{path:path}")
