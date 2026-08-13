"""The coverage signal (spec T6).

Fixtures are built in-process with PyMuPDF and python-docx rather than
committed, so this suite opens no corpus and loads no model.
"""
from pathlib import Path

import docx
import fitz
import pytest

from ingest.coverage import COVERAGE_FLOOR, coverage_ratio, source_text_chars


def _pdf(tmp_path: Path, pages: list[str], name: str = "f.pdf") -> Path:
    doc = fitz.open()
    for body in pages:
        page = doc.new_page()
        if body:
            page.insert_text((72, 72), body)
    path = tmp_path / name
    doc.save(path)
    doc.close()
    return path


def _docx(
    tmp_path: Path,
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
    name: str = "f.docx",
) -> Path:
    document = docx.Document()
    for text in paragraphs or []:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row_vals in enumerate(table_rows):
            for c, val in enumerate(row_vals):
                table.cell(r, c).text = val
    path = tmp_path / name
    document.save(path)
    return path


def test_source_text_chars_counts_the_text_layer(tmp_path):
    path = _pdf(tmp_path, ["hello world", "second page"])
    assert source_text_chars(path) >= len("hello world") + len("second page")


def test_ratio_is_produced_over_source(tmp_path):
    path = _pdf(tmp_path, ["abcdefghij"])          # 10 chars of source text
    got = coverage_ratio(["abcde"], path)          # 5 chars produced
    assert got == pytest.approx(5 / source_text_chars(path))


def test_the_fy2024_afr_shape_lands_below_the_floor(tmp_path):
    """20 chunks from a 191-page document scored 2.0% on the real corpus."""
    path = _pdf(tmp_path, ["x" * 1000 for _ in range(10)])
    ratio = coverage_ratio(["x" * 20], path)
    assert ratio < COVERAGE_FLOOR


def test_a_ratio_above_one_is_returned_unchanged(tmp_path):
    """Healthy AFRs score 278-286% because chunk text carries table markup.

    Clamping to 1.0 would erase the single clearest signal that extraction
    is working. Pinned because "normalize it to a percentage" is a natural
    and wrong instinct.
    """
    path = _pdf(tmp_path, ["abc"])
    assert coverage_ratio(["x" * 10_000], path) > 1.0


def test_no_text_layer_returns_None_rather_than_zero(tmp_path):
    """An image-only PDF must route to OCR, not read as a failed extraction.

    0.0 and None are different answers: 0.0 means "we extracted nothing from
    a document that has text", None means "there is nothing here to compare
    against".
    """
    path = _pdf(tmp_path, [""])
    assert coverage_ratio(["anything"], path) is None


def test_the_floor_is_the_calibrated_value():
    """Pinned so a future edit has to go and re-read the calibration."""
    assert COVERAGE_FLOOR == 0.10


def test_source_text_chars_docx_includes_both_paragraphs_and_table_cells(tmp_path):
    """Budget bills are DOCX-only by design, so this path takes a whole
    document class. A paragraphs-only sum would undercount any DOCX that
    carries text in a table, silently -- no error, just a smaller number.
    """
    paragraphs = ["Section 1. Appropriations.", "This act amends the budget."]
    table_rows = [["Agency", "Amount"], ["AHCCCS", "1,234,567,890"]]
    path = _docx(tmp_path, paragraphs=paragraphs, table_rows=table_rows)

    paragraph_chars = sum(len(p) for p in paragraphs)
    cell_chars = sum(len(v) for row in table_rows for v in row)

    got = source_text_chars(path)

    assert got >= paragraph_chars + cell_chars
    # A paragraphs-only implementation would stop here and be too small --
    # the table-cell text has to have been added on top.
    assert got > paragraph_chars


def test_source_text_chars_docx_counts_table_only_text(tmp_path):
    """A DOCX whose text lives entirely in table cells, no paragraphs.

    This is the test that actually pins the module's WHY comment: a
    paragraphs-only implementation returns 0 here, which would make every
    table-shaped DOCX -- exactly the budget-bill shape this module was
    written for -- read as a total extraction failure.
    """
    table_rows = [["AHCCCS", "1,234,567,890"], ["ADE", "9,876,543,210"]]
    path = _docx(tmp_path, paragraphs=[], table_rows=table_rows)

    assert source_text_chars(path) > 0


def test_source_text_chars_raises_on_an_unsupported_suffix(tmp_path):
    path = tmp_path / "f.txt"
    path.write_text("hello", encoding="utf-8")

    with pytest.raises(ValueError):
        source_text_chars(path)
