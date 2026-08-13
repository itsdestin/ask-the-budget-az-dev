"""Look at a source file and decide where to START extracting (spec T5 step 1).

## Inspection picks the starting rung. It NEVER decides success.

This is the decisive measurement behind the whole design and it is worth
stating at the code: `agao-afr-fy2023` (healthy, 281% coverage) and
`agao-afr-fy2024` (broken, 2.0%) are INDISTINGUISHABLE here. Both report
191/184 pages and ~1.1M characters of text layer. Anything that tries to
predict failure from inspection alone will pass the FY2024 AFR, which is the
document this design exists for. Only running the extraction and measuring
its output (T6) separates them.

## Why there is no tagging field, despite spec T7 asking for one

Measured across every OpenDataLoader-first document in the corpus:

    governor-governors-budget-fy2025   UNTAGGED, 639pp   ->  92.2%
    governor-governors-budget-fy2026   tagged,   661pp   ->  96.0%
    agao-afr-fy2024                    TAGGED,   191pp   ->   2.0%

Tagging does not predict success. OpenDataLoader reads a large untagged PDF
fine, and the one document that fails is tagged. A "no structure tree ->
skip OpenDataLoader" rule would therefore divert a healthy 639-page document
to a slower extractor and change its chunk text for nothing. Do not add the
field back.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class SourceInspection:
    """What the file itself will tell us before anything is extracted."""

    source_format: str
    # None for DOCX (no pages at rest) and for a file we could not open.
    # None rather than 0: "not applicable" and "empty" are different.
    pages: int | None
    has_text_layer: bool


def inspect_source(path: Path) -> SourceInspection:
    """Cheap, total, and never raises.

    An inspection failure is a valid answer -- it means we learned nothing
    and the ladder starts at rung 1. Raising here would take down the worker
    thread over a truncated download.
    """
    source_format = path.suffix.lstrip(".").lower()

    if source_format == "pdf":
        try:
            import fitz  # PyMuPDF

            with fitz.open(path) as doc:
                # `.strip()` matters: a PDF of scanned pages often carries a
                # few whitespace glyphs, which would otherwise read as text
                # and route a scan away from the OCR rung it needs.
                has_text = any(page.get_text().strip() for page in doc)
                return SourceInspection("pdf", doc.page_count, has_text)
        except Exception:
            return SourceInspection("pdf", None, False)

    if source_format == "docx":
        try:
            import docx  # python-docx

            document = docx.Document(str(path))
            has_text = any(p.text.strip() for p in document.paragraphs) or any(
                cell.text.strip()
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            )
            return SourceInspection("docx", None, has_text)
        except Exception:
            return SourceInspection("docx", None, False)

    return SourceInspection(source_format, None, False)
