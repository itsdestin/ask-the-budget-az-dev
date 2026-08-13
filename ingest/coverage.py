"""Does an extraction look catastrophically empty? (spec T6)

The signal is characters of chunk text produced divided by characters in the
source file's own text layer. It exists because `agao-afr-fy2024` reported
`live` with 20 passages from 191 pages, the queue showed green, and an
analyst searching FY2024 simply got nothing.

## What this does NOT do

It detects catastrophic LOSS, not CORRUPTION. A document that produced the
right amount of the wrong text passes. That is not hypothetical on this
corpus: the FY2024 AFR's own recovered rows are label-stripped table
fragments, and a numeric-density check scored them 1.6% "junk" -- apparently
clean -- because they are full of agency and fund names. Passing this check
is not a certificate of health, and nothing may describe it as one.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

# CALIBRATED 2026-08-12 across all 7,434 documents, post orphan-recovery
# repair. Full measurement:
#   docs/superpowers/investigations/2026-08-12-coverage-floor-calibration.md
#
# Median coverage is 87.9%. Every floor from just above 2.0% to just below
# 17.1% catches an IDENTICAL set of two documents, so this is a plateau and
# 0.10 is its centre -- the right pick because the metric degrades on both
# sides: below 2.0% the known-broken AFR escapes, above 17.1% healthy short
# documents start being caught.
#
# The spec's original expectation was 15-25%, from a 16-document sample taken
# before the orphan-recovery bug was fixed. The corpus-wide run says that is
# too high. Do not restore it without re-running the measurement.
COVERAGE_FLOOR = 0.10


def source_text_chars(path: Path) -> int:
    """Characters in the source file's own text layer -- the denominator.

    This reads the SOURCE, deliberately, not the extractor's output: the
    question is "how much of what is in this file came out", and only the
    file itself can answer it.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        import fitz  # PyMuPDF

        with fitz.open(path) as doc:
            return sum(len(page.get_text()) for page in doc)
    if suffix == ".docx":
        import docx  # python-docx

        document = docx.Document(str(path))
        # Table cells are not in `paragraphs` and a budget bill is mostly
        # tables, so counting paragraphs alone would make every DOCX look
        # like a failed extraction.
        body = sum(len(p.text) for p in document.paragraphs)
        cells = sum(
            len(cell.text)
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        return body + cells
    raise ValueError(f"coverage: no text-layer reader for {suffix!r} ({path.name})")


def coverage_ratio(chunk_texts: Iterable[str], source_path: Path) -> float | None:
    """Produced characters over source characters.

    Returns None when the source has no text layer at all. That is a ROUTING
    signal -- an image-only PDF goes to OCR -- and must not be confused with
    0.0, which means "this document has text and we extracted none of it".

    The result is NOT clamped. Values above 1.0 are normal and are the
    clearest evidence extraction is working; see the module docstring.
    """
    produced = sum(len(text or "") for text in chunk_texts)
    total = source_text_chars(source_path)
    if total == 0:
        return None
    return produced / total
