"""`create_document` — the model's only write primitive (Plan 4 Task 4, S3).

An analyst asks for a memo-shaped answer ("write this up so I can send
it"); the model calls `create_document` with a title and a Markdown
body; this module turns that into a .docx (or .md) the analyst can
download, and hands back an unguessable token the chat UI renders as a
link.

**Invariant 7 lives here.** One Python process serves a whole office off
a shared network drive — the corpus, settings and PDFs all live on that
share. Nothing the model can call may write there, because a confused or
prompt-injected model would then be able to corrupt state for everyone
at once. So:

* The tool schema takes a TITLE, never a path or a filename. The model
  supplies content; this module supplies the location. That split is the
  whole invariant.
* Artifacts land under the per-user, per-machine local app-data folder
  (`%LOCALAPPDATA%\\JLBC-Insight\\documents\\`) — disposable, private,
  and not the share.
* This module deliberately does NOT import `store.config`, so it has no
  way to learn where the share even is. `tests/test_create_document.py`
  asserts that as an import allowlist, which is why the guard is
  structural rather than a promise in a comment.

Kept import-light on purpose: `harness/tools.py` imports it lazily
inside its `create_document` handler, so anything expensive at module
scope would be paid on the first tool dispatch of a conversation.
python-docx (and its lxml dependency) is therefore imported inside the
writer, not here — the .md path never pays for it at all.
"""
from __future__ import annotations

import os
import re
import secrets
import threading
import time
import unicodedata
from dataclasses import dataclass
from pathlib import Path

# Override for tests, and an escape hatch for an install that wants
# artifacts somewhere else. Named for what it holds so it can never be
# confused with JLBC_DATA_DIR, which points at the SHARED data root —
# these two must never be the same directory (see documents_dir()).
DOCUMENTS_DIR_ENV = "JLBC_DOCUMENTS_DIR"

_APP_FOLDER = "JLBC-Insight"

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# fmt -> (file extension, HTTP content type). The single source of truth
# for both; the download route reads the media type off the registry
# rather than re-deriving it from the suffix.
_FORMATS: dict[str, tuple[str, str]] = {
    "docx": (".docx", DOCX_MEDIA_TYPE),
    "md": (".md", "text/markdown; charset=utf-8"),
}


def documents_dir() -> Path:
    """Where artifacts are written. Created if missing.

    Resolution order:
      1. `JLBC_DOCUMENTS_DIR` — tests point this at a tmp dir.
      2. `%LOCALAPPDATA%\\JLBC-Insight\\documents` on Windows, which is
         the real deployment (spec S7 installs the whole app there).
      3. Non-Windows (CI, a dev Mac) has no LOCALAPPDATA, so fall back to
         the XDG data location — `~/.local/share/JLBC-Insight/documents`.
         Home-rooted either way, which is what matters: never the share.

    NOT validated against the shared data dir, because doing so would
    mean importing `store.config` and giving this module knowledge of the
    share it is specifically supposed to lack. Pointing
    JLBC_DOCUMENTS_DIR at the share would defeat Invariant 7, but that is
    an operator typing a deliberate env var, not something any model can
    reach through the tool surface — and the tool surface is what the
    invariant is about.
    """
    raw = os.environ.get(DOCUMENTS_DIR_ENV)
    if raw:
        root = Path(raw)
    elif os.environ.get("LOCALAPPDATA"):
        root = Path(os.environ["LOCALAPPDATA"]) / _APP_FOLDER / "documents"
    else:
        base = os.environ.get("XDG_DATA_HOME") or (Path.home() / ".local" / "share")
        root = Path(base) / _APP_FOLDER / "documents"
    root.mkdir(parents=True, exist_ok=True)
    return root


# ---------------------------------------------------------------------------
# The download registry
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Artifact:
    """One materialized document, as the download route needs it."""

    token: str
    path: Path
    media_type: str
    # WHO asked for it. Recorded for attribution and debugging, NOT
    # enforced at download time — and that is deliberate rather than an
    # oversight. This app has no authentication (spec S11 is explicit
    # that even the admin surface is soft-gated and "explicitly not real
    # security"), so a username is a string the caller asserts about
    # itself. Checking it would look like isolation while providing
    # none. The real protection is that the token is 256 bits of
    # `secrets` randomness that only ever travels to the conversation
    # that created it. If per-user isolation ever has to be REAL, it
    # needs an authenticated identity first; bolting a comparison onto
    # this field would just be theater.
    user: str
    created_at: float


_registry_lock = threading.Lock()
_registry: dict[str, Artifact] = {}


def reset_registry() -> None:
    """Test-only: drop every token. Does not delete files from disk."""
    with _registry_lock:
        _registry.clear()


def lookup(token: str) -> Artifact | None:
    """Resolve a download token, or None if it is unknown.

    THE TOKEN IS A DICTIONARY KEY AND NEVER A PATH COMPONENT. That is
    what makes path traversal structurally impossible rather than merely
    filtered: a token of `../../settings.json` is not a dangerous path,
    it is an absent key. Nothing a caller sends is ever joined onto a
    directory.

    The registry is in-process, so tokens die with the server (accepted
    per the plan — conversation persistence across restarts is a Plan 5
    item). The FILES outlive it, orphaned under documents_dir(); they are
    small, local and disposable, and sweeping them is a launcher job for
    Plan 5 rather than something to invent here.
    """
    with _registry_lock:
        return _registry.get(token)


# ---------------------------------------------------------------------------
# Filenames
# ---------------------------------------------------------------------------
# The model chooses the title, so the title is untrusted input that ends
# up in two hostile places: a filesystem path and an HTTP header. This is
# the only point where a model-supplied string gets anywhere near either.

# Kept deliberately tighter than "what NTFS accepts": exactly the
# characters `urllib.parse.quote` leaves untouched, so an ASCII title
# always yields the simple `filename="…"` Content-Disposition form rather
# than an RFC 5987 escape, and so `/ \ : * ? " < > |`, control characters
# and CR/LF (header injection) cannot survive by construction — they are
# not on the list, so they are replaced rather than looked for.
_KEEP = set("-_.")
# Windows refuses these names with ANY extension — `NUL.docx` is still the
# null device, and writing to it silently succeeds while producing no file.
_RESERVED = {"CON", "PRN", "AUX", "NUL"} | {
    f"{prefix}{n}" for prefix in ("COM", "LPT") for n in range(1, 10)
}
# Long enough for a real memo title, short enough that stem + extension +
# the per-artifact folder stay well inside Windows' 260-char path limit
# on a deep %LOCALAPPDATA% path.
_MAX_STEM = 80


def _safe_stem(title: str) -> str:
    """Untrusted title -> a filename stem that is safe everywhere."""
    text = unicodedata.normalize("NFC", title)
    # Allowlist, not denylist: anything unrecognized becomes a space and
    # is then collapsed, so no exotic separator or control character has
    # to be individually anticipated.
    kept = "".join(ch if (ch.isalnum() or ch in _KEEP) else " " for ch in text)
    stem = "-".join(kept.split())
    stem = re.sub(r"-{2,}", "-", stem)
    # Windows rejects trailing dots/spaces, and a leading dot would make
    # the file hidden on POSIX.
    stem = stem.strip("-. ")[:_MAX_STEM].strip("-. ")
    if not stem:
        return "document"  # a title of "..." or "***" leaves nothing
    if stem.upper() in _RESERVED or stem.split(".")[0].upper() in _RESERVED:
        return f"document-{stem}"
    return stem


# ---------------------------------------------------------------------------
# Markdown -> Word
# ---------------------------------------------------------------------------
# A small, deliberate subset — the inverse of what primer/docx_to_md.py
# emits, which is also what the system prompt tells the model to write.
#
# THE RULE THAT MATTERS: anything unrecognized becomes a plain paragraph,
# verbatim. Never a silent drop. An analyst who receives a memo with a
# section quietly missing has no way to know it happened, and that is a
# far worse failure than a blockquote rendering as ordinary text.

_HEADING_RE = re.compile(r"^(#{1,6})\s+(\S.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(\S.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
# A separator row: only dashes, colons, pipes and spaces, and at least
# one dash. Requiring the PRECEDING line to be a row too (see the loop)
# keeps a bare `---` thematic break from being mistaken for a table.
_TABLE_SEPARATOR_RE = re.compile(r"^\|?[\s:|-]*-[\s:|-]*\|?$")
# Split on pipes that are not backslash-escaped — docx_to_md.py escapes a
# literal pipe inside a cell as `\|`, so this is the matching unescape.
_UNESCAPED_PIPE_RE = re.compile(r"(?<!\\)\|")


def _is_table_row(line: str) -> bool:
    return "|" in line


def _split_row(line: str) -> list[str]:
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|") and not inner.endswith("\\|"):
        inner = inner[:-1]
    return [c.strip().replace("\\|", "|") for c in _UNESCAPED_PIPE_RE.split(inner)]


def _add_runs(paragraph, text: str) -> None:
    """Write `text` into a paragraph, turning **…** into bold runs.

    Everything outside the markers is emitted as-is, so an unmatched `**`
    stays visible rather than eating the rest of the line.
    """
    position = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _add_table(doc, rows: list[list[str]]) -> None:
    """Render collected pipe-table rows as a real Word table.

    Short rows are PADDED rather than dropped, matching what
    primer/docx_to_md.py does in the other direction: a ragged row is a
    formatting slip in the model's output, not a reason to lose the cell
    values it does contain.
    """
    columns = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(columns):
            value = row[column] if column < len(row) else ""
            paragraph = cells[column].paragraphs[0]
            _add_runs(paragraph, value)
            if index == 0:
                # Header row: bold every run, including ones **…** already
                # bolded (idempotent).
                for run in paragraph.runs:
                    run.bold = True


def _render_docx(title: str, body_markdown: str, target: Path) -> None:
    """Write the .docx. python-docx is imported HERE, not at module
    scope — see the module docstring on staying import-light."""
    from docx import Document

    doc = Document()
    doc.core_properties.title = title
    # The title leads the document as a Word Title-styled heading. The
    # filename alone would leave the artifact untitled once it is printed
    # or pasted into an email.
    _add_runs(doc.add_heading("", level=0), title)

    lines = body_markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        # Tables first: a table row is only a table row when the NEXT
        # line is a separator, otherwise it is ordinary text containing
        # pipe characters and must survive as such.
        if (
            _is_table_row(stripped)
            and index + 1 < len(lines)
            and _TABLE_SEPARATOR_RE.match(lines[index + 1].strip())
            and "|" in lines[index + 1]
        ):
            rows = [_split_row(stripped)]
            index += 2  # header + separator
            while index < len(lines) and _is_table_row(lines[index].strip()):
                rows.append(_split_row(lines[index].strip()))
                index += 1
            _add_table(doc, rows)
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = len(heading.group(1))
            _add_runs(doc.add_heading("", level=level), heading.group(2))
            index += 1
            continue

        bullet = _BULLET_RE.match(stripped)
        if bullet:
            _add_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            index += 1
            continue

        # Everything else — blockquotes, numbered lists, code fences,
        # links, tables missing a separator — lands here VERBATIM. The
        # markup is visible but no content is lost.
        _add_runs(doc.add_paragraph(), line)
        index += 1

    doc.save(str(target))


# ---------------------------------------------------------------------------
# The public entry point
# ---------------------------------------------------------------------------


def materialize(
    title: str, body_markdown: str, fmt: str = "docx", *, user: str = ""
) -> tuple[str, Path]:
    """Write one artifact and register it for download.

    Returns `(token, path)`. `harness/tools.py` reads `path.name` as the
    filename it reports to the model, which is why the path keeps the
    human title rather than being named after the token.

    Each artifact gets its OWN randomly-named subdirectory. Two memos
    titled "Memo" would otherwise collide, and the second write would
    silently replace the first while the first token still pointed at
    that path — handing an analyst someone else's document under their
    own filename.
    """
    if fmt not in _FORMATS:
        raise ValueError(
            f"unknown format {fmt!r} — must be one of: {', '.join(_FORMATS)}."
        )
    suffix, media_type = _FORMATS[fmt]

    folder = documents_dir() / secrets.token_hex(8)
    folder.mkdir(parents=True, exist_ok=True)
    target = folder / f"{_safe_stem(title)}{suffix}"
    # Belt and braces over _safe_stem: assert the sanitizer's output
    # actually landed where it was supposed to, so a future edit that
    # loosens the allowlist fails loudly instead of writing outside the
    # documents directory.
    if target.resolve().parent != folder.resolve():
        raise ValueError(f"refusing to write outside {folder}")

    if fmt == "md":
        # Byte-faithful: the analyst downloads exactly what the model
        # wrote, no round-trip through a renderer that could lose a
        # construct.
        target.write_text(body_markdown, encoding="utf-8")
    else:
        _render_docx(title, body_markdown, target)

    # 32 bytes -> a 43-character URL-safe string. This token is the ONLY
    # thing standing between an artifact and anyone who can reach the
    # loopback port, so it is generated by `secrets` and is unrelated to
    # the title, the user, or the path.
    token = secrets.token_urlsafe(32)
    with _registry_lock:
        _registry[token] = Artifact(
            token=token,
            path=target,
            media_type=media_type,
            user=user,
            created_at=time.time(),
        )
    return token, target
